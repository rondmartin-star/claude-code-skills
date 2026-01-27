#!/usr/bin/env python3
"""
Convert HTML content to DOCX format.

Usage:
    python convert_to_docx.py <input.html> [--output <output.docx>] [--template <template.docx>]

Requirements:
    pip install python-docx beautifulsoup4 lxml

Example:
    python convert_to_docx.py article.html
    python convert_to_docx.py article.html --output submission.docx
"""

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from bs4 import BeautifulSoup
    HAS_DEPENDENCIES = True
except ImportError:
    HAS_DEPENDENCIES = False


class HTMLToDocxConverter:
    """Convert HTML content to DOCX format."""
    
    def __init__(self, template_path: Path = None):
        if template_path and template_path.exists():
            self.doc = Document(str(template_path))
        else:
            self.doc = Document()
            self._setup_styles()
    
    def _setup_styles(self):
        """Set up default document styles."""
        # Set default font
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Set paragraph formatting
        para_format = style.paragraph_format
        para_format.line_spacing = 2.0  # Double spacing
        para_format.space_after = Pt(0)
        
        # Set margins
        for section in self.doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
    
    def convert(self, html_content: str) -> Document:
        """Convert HTML to DOCX."""
        soup = BeautifulSoup(html_content, 'lxml')
        
        # Remove script and style tags
        for tag in soup.find_all(['script', 'style', 'div.controls']):
            tag.decompose()
        
        # Find main content
        main = soup.find('main') or soup.find('article') or soup.find('body')
        if not main:
            main = soup
        
        self._process_element(main)
        
        return self.doc
    
    def _process_element(self, element):
        """Recursively process HTML elements."""
        if element.name is None:
            return
        
        for child in element.children:
            if hasattr(child, 'name'):
                self._handle_tag(child)
            elif isinstance(child, str) and child.strip():
                # Text node outside of tags
                pass
    
    def _handle_tag(self, tag):
        """Handle individual HTML tags."""
        tag_name = tag.name.lower() if tag.name else ''
        
        if tag_name == 'h1':
            self._add_heading(tag, level=0)
        elif tag_name == 'h2':
            self._add_heading(tag, level=1)
        elif tag_name == 'h3':
            self._add_heading(tag, level=2)
        elif tag_name == 'h4':
            self._add_heading(tag, level=3)
        elif tag_name == 'p':
            self._add_paragraph(tag)
        elif tag_name == 'ul':
            self._add_list(tag, ordered=False)
        elif tag_name == 'ol':
            self._add_list(tag, ordered=True)
        elif tag_name == 'blockquote':
            self._add_blockquote(tag)
        elif tag_name == 'table':
            self._add_table(tag)
        elif tag_name == 'figure':
            self._add_figure(tag)
        elif tag_name == 'img':
            self._add_image(tag)
        elif tag_name == 'hr':
            self._add_horizontal_rule()
        elif tag_name in ['div', 'section', 'article', 'header', 'footer', 'main']:
            # Container elements - process children
            self._process_element(tag)
    
    def _get_text(self, element) -> str:
        """Get text content from element."""
        return element.get_text(strip=True)
    
    def _add_heading(self, tag, level: int):
        """Add a heading to the document."""
        text = self._get_text(tag)
        if text:
            heading = self.doc.add_heading(text, level=level)
            # Center title (level 0)
            if level == 0:
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_paragraph(self, tag):
        """Add a paragraph to the document."""
        text = self._get_text(tag)
        if text:
            para = self.doc.add_paragraph()
            
            # Handle mixed content (text with bold/italic)
            self._add_formatted_text(para, tag)
            
            # Add first-line indent for regular paragraphs
            para.paragraph_format.first_line_indent = Inches(0.5)
    
    def _add_formatted_text(self, paragraph, element):
        """Add text with formatting (bold, italic, links)."""
        for child in element.children:
            if isinstance(child, str):
                if child.strip():
                    paragraph.add_run(child)
            elif child.name == 'strong' or child.name == 'b':
                run = paragraph.add_run(self._get_text(child))
                run.bold = True
            elif child.name == 'em' or child.name == 'i':
                run = paragraph.add_run(self._get_text(child))
                run.italic = True
            elif child.name == 'a':
                run = paragraph.add_run(self._get_text(child))
                run.underline = True
                run.font.color.rgb = RGBColor(0, 0, 255)
            elif child.name == 'sup':
                run = paragraph.add_run(self._get_text(child))
                run.font.superscript = True
            elif child.name == 'sub':
                run = paragraph.add_run(self._get_text(child))
                run.font.subscript = True
            elif child.name == 'br':
                paragraph.add_run('\n')
            elif child.name == 'span':
                # Check for citation class
                if 'citation' in child.get('class', []):
                    run = paragraph.add_run(self._get_text(child))
                    run.font.color.rgb = RGBColor(0, 0, 255)
                else:
                    self._add_formatted_text(paragraph, child)
            else:
                # Recursively handle nested elements
                self._add_formatted_text(paragraph, child)
    
    def _add_list(self, tag, ordered: bool):
        """Add a list to the document."""
        style = 'List Number' if ordered else 'List Bullet'
        
        for li in tag.find_all('li', recursive=False):
            para = self.doc.add_paragraph(style=style)
            self._add_formatted_text(para, li)
    
    def _add_blockquote(self, tag):
        """Add a blockquote to the document."""
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.first_line_indent = Inches(0)
        
        run = para.add_run(self._get_text(tag))
        run.italic = True
    
    def _add_table(self, tag):
        """Add a table to the document."""
        # Find all rows
        rows = tag.find_all('tr')
        if not rows:
            return
        
        # Determine number of columns from first row
        first_row = rows[0]
        cols = len(first_row.find_all(['th', 'td']))
        
        if cols == 0:
            return
        
        # Create table
        table = self.doc.add_table(rows=len(rows), cols=cols)
        table.style = 'Table Grid'
        
        for i, row in enumerate(rows):
            cells = row.find_all(['th', 'td'])
            for j, cell in enumerate(cells):
                if j < cols:
                    table_cell = table.rows[i].cells[j]
                    table_cell.text = self._get_text(cell)
                    
                    # Bold header cells
                    if cell.name == 'th':
                        for para in table_cell.paragraphs:
                            for run in para.runs:
                                run.bold = True
        
        # Add space after table
        self.doc.add_paragraph()
    
    def _add_figure(self, tag):
        """Add a figure with caption."""
        img = tag.find('img')
        caption = tag.find('figcaption')
        
        if img:
            self._add_image(img)
        
        if caption:
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(self._get_text(caption))
            run.italic = True
            run.font.size = Pt(10)
    
    def _add_image(self, tag):
        """Add an image to the document."""
        src = tag.get('src', '')
        alt = tag.get('alt', '')
        
        # Try to add image if it's a local file
        if src and not src.startswith(('http://', 'https://', 'data:')):
            try:
                para = self.doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(src, width=Inches(5))
            except Exception as e:
                # If image can't be added, add placeholder text
                para = self.doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.add_run(f"[Image: {alt or src}]")
        else:
            # External image - add placeholder
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run(f"[Image: {alt or src}]")
    
    def _add_horizontal_rule(self):
        """Add a horizontal rule (as paragraph with border)."""
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(12)
        # Note: python-docx doesn't have direct HR support, so we add empty space
    
    def save(self, output_path: Path):
        """Save the document."""
        self.doc.save(str(output_path))


def main():
    parser = argparse.ArgumentParser(description="Convert HTML to DOCX")
    parser.add_argument("input", help="Input HTML file")
    parser.add_argument("--output", "-o", help="Output DOCX file")
    parser.add_argument("--template", "-t", help="Template DOCX file")
    
    args = parser.parse_args()
    
    if not HAS_DEPENDENCIES:
        print("Error: Required dependencies not installed.")
        print("Run: pip install python-docx beautifulsoup4 lxml")
        return 1
    
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: File not found: {input_path}")
        return 1
    
    output_path = Path(args.output) if args.output else input_path.with_suffix('.docx')
    template_path = Path(args.template) if args.template else None
    
    # Read HTML
    html_content = input_path.read_text(encoding='utf-8')
    
    # Convert
    converter = HTMLToDocxConverter(template_path)
    converter.convert(html_content)
    converter.save(output_path)
    
    print(f"✓ Created: {output_path}")
    return 0


if __name__ == "__main__":
    exit(main())
